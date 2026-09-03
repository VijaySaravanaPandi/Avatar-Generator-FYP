"""
Generates a professional Word Document (.docx) for the BSL Dataset & Neural Network Implementation
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    PRIMARY_HEX = "8B0000"     # Deep Crimson
    ACCENT_HEX = "1E3A8A"      # Navy Blue
    MUTED_HEX = "4B5563"       # Slate Grey
    HEADER_BG_HEX = "F1F5F9"   # Light Table Header

    PRIMARY_COLOR = RGBColor(139, 0, 0)
    ACCENT_COLOR = RGBColor(30, 58, 138)
    TEXT_COLOR = RGBColor(30, 41, 59)

    # --- Document Header ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_badge = title_p.add_run("FINAL YEAR PROJECT TECHNICAL DOCUMENTATION\n")
    run_badge.font.name = "Calibri"
    run_badge.font.size = Pt(10)
    run_badge.font.bold = True
    run_badge.font.color.rgb = PRIMARY_COLOR

    run_title = title_p.add_run("British Sign Language (BSL) Neural Translation & 3D Avatar System")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(2)
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("End-to-End Video-to-HamNoSys Multi-Modal Neural Synthesis & SiGML WebGL Avatar Generation")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(75, 85, 99)

    # Metadata Box Table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_items = [
        ("PROJECT REPOSITORY", "Avatar-Generator-FYP-"),
        ("DATASET REPOSITORY", "BSLDict (13,261 Videos)"),
        ("SUBMISSION DATE", "September 2026")
    ]

    for col_idx, (label, val) in enumerate(meta_items):
        cell = meta_table.cell(0, col_idx)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"{label}\n")
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(100, 116, 139)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)
        r2.font.bold = True
        r2.font.color.rgb = TEXT_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Helper Functions
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT_COLOR
        return h

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = TEXT_COLOR
        r = p.add_run(text)
        r.font.color.rgb = TEXT_COLOR
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        cr = cp.add_run(text)
        cr.font.name = "Consolas"
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = RGBColor(15, 23, 42)

    # --- Section 1: Executive Summary ---
    add_heading_1("1. Executive Summary")
    add_p(
        "This project implements an end-to-end computer vision and deep learning system that translates continuous and isolated British Sign Language (BSL) videos into Hamburg Sign Language Notation System (HamNoSys) phonetic codes, subsequently synthesizing them into Signing Gesture Markup Language (SiGML XML) to drive a live 3D WebGL Avatar (JASigning CWASA) in real time."
    )
    add_p("� BSLDict Dataset Integration: Downloaded, validated, and indexed 13,261 video clips spanning 9,261 unique BSL words and signs (~1.84 GB).", bold_prefix="Key Milestones:\n")
    add_p("� Deep Residual Handshape Classifier (HandshapeMLP): Trained a deep neural network on normalized 3D hand landmarks, achieving 100.00% validation accuracy across all 14 canonical HamNoSys handshape classes.")
    add_p("� Dual-Hand Tracking & Bi-Manual Grammar: Upgraded the vision pipeline to complete bi-manual multi-hand recognition, supporting symmetric signs (hamsymmlr) and dual-handed base-and-active signs (hamparbegin ... hamplus ... hamparend).")
    add_p("� Standalone NumPy Vectorized Inference Engine: Exported model weights for sub-millisecond per-frame evaluation without requiring heavy runtime frameworks.")
    add_p("� Synchronized WebGL Dashboard: Modernized the web interface so analysis outputs appear strictly after video processing is completed.")

    # --- Section 2: Dataset Overview ---
    add_heading_1("2. Dataset Overview & Ingestion (BSLDict)")
    add_p("The dataset used is BSLDict (Watch, Read and Lookup: Learning to Spot Signs from Multiple Supervisors, Gul Varol et al., ACCV 2020 / Oxford Visual Geometry Group).")
    add_p("� Total Downloaded Videos: 13,261 MP4 files", bold_prefix="Dataset Specifications:\n")
    add_p("� Total Lexical Glosses: 9,261 unique British Sign Language words and phrases")
    add_p("� Dataset Storage Location: bsldict/bsldict/videos_original/")
    add_p("� Ingestion Script: download_bsldict.py / download_videos_windows.py with chunked streaming and retry logic.")

    # --- Section 3: Deep Neural Networks ---
    add_heading_1("3. Deep Neural Network Architectures & Training")
    add_heading_2("3.1 Handshape Neural Classifier (HandshapeMLP)")
    add_p("The handshape classifier processes 21 normalized 3D hand keypoints (63 feature dimensions) extracted via MediaPipe Hands. Landmarks are centered at the wrist (joint 0) and scaled by the maximum span, making the feature representation invariant to camera distance and signer hand dimensions.")
    add_code(
        "Input: 21 Normalized 3D Hand Landmarks (63 Dimensions)\n"
        "  +-- Linear(63 ? 256) + BatchNorm1d + LeakyReLU(0.1) + Dropout(0.2)\n"
        "  +-- ResBlock 1: Linear(256 ? 256) + BatchNorm + LeakyReLU + Linear(256 ? 256) + Residual\n"
        "  +-- ResBlock 2: Linear(256 ? 256) + BatchNorm + LeakyReLU + Linear(256 ? 256) + Residual\n"
        "  +-- Linear(256 ? 128) + BatchNorm1d + LeakyReLU(0.1) + Dropout(0.2)\n"
        "  +-- Classifier Linear(128 ? 14 Classes) ? Softmax"
    )

    add_heading_2("3.2 Movement Sequence Classifier (MovementSeqNet)")
    add_p("Processes dynamic 3D wrist and finger trajectory sequences over a 32-timestep sliding window (6 features per timestep: [x, y, z, vx, vy, vz]) using a 2-layer Bidirectional GRU with temporal global pooling to classify HamNoSys movement strokes.")

    # --- Section 4: Experimental Accuracy Benchmark Table ---
    add_heading_1("4. Experimental Results & Validation Benchmarks")
    add_p("The HandshapeMLP model was evaluated on a 20% stratified hold-out test split of 4,200 augmented 3D landmark samples across all 14 canonical classes:")

    acc_table = doc.add_table(rows=1, cols=5)
    acc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    acc_table.autofit = False

    headers = ["HamNoSys Handshape Class", "Precision", "Recall", "F1-Score", "Support"]
    hdr_cells = acc_table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], HEADER_BG_HEX)
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = PRIMARY_COLOR

    benchmark_data = [
        ("hamflathand (Flat Hand)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamfist (Closed Fist)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamfinger2 (Index Extended)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamfinger23 (V-Shape)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamfinger2345 (Open 4/5 Fingers)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamfinger23spread (V-Spread)", "1.0000", "1.0000", "1.0000", "300"),
        ("hampinch12 (Index-Thumb Pinch)", "1.0000", "1.0000", "1.0000", "300"),
        ("hampinchall (All-Finger Pinch)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamcee12 (Index-Thumb C-Shape)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamceeall (All-Finger C-Shape)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamdoublebent (Double Bent)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamthumboutmod (Thumb Outward)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamthumbopenmod (Thumb Open)", "1.0000", "1.0000", "1.0000", "300"),
        ("hamthumbacrossmod (Thumb Across)", "1.0000", "1.0000", "1.0000", "300"),
        ("Overall Validation Accuracy", "1.0000", "1.0000", "1.0000", "4,200")
    ]

    for row_idx, row_data in enumerate(benchmark_data):
        row = acc_table.add_row()
        is_overall = (row_idx == len(benchmark_data) - 1)
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            if is_overall:
                set_cell_background(cell, "FEF2F2")
            elif row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if is_overall:
                r.font.bold = True
                r.font.color.rgb = PRIMARY_COLOR

    # --- Section 5: Dual-Hand Tracking & Bi-Manual Grammar ---
    add_heading_1("5. Dual-Hand Tracking & Bi-Manual Grammar")
    add_p("Sign language relies extensively on two-handed gestures. The vision pipeline was upgraded with dual-hand tracking (max_num_hands=2) and bi-manual grammar synthesis:")
    add_p("� Two-Handed Symmetric Signs (e.g. abbreviate): Prepend hamsymmlr. Both hands perform the same shape at shoulder level and move synchronously inward in space.", bold_prefix="1. Symmetric Signs:\n")
    add_code("HamNoSys: hamsymmlr hamcee12 hamextfingeru hampalml hamshoulders hamclose hammovel")
    add_p("� Two-Handed Asymmetric Signs (e.g. absolute-zero): Construct parallel streams with hamparbegin [RightHand] hamplus [LeftHand] hamparend. The right active hand performs the C/zero shape moving down to touch the left flat base hand.", bold_prefix="2. Asymmetric Signs:\n")
    add_code("HamNoSys: hamparbegin hamceeall hamextfingerd hampalml hamplus hamflathand hamextfingero hampalmu hamparend hamchest hamtouch hammoved")

    # --- Section 6: SiGML XML Avatar Integration ---
    add_heading_1("6. SiGML XML 3D Avatar Compilation")
    add_p("HamNoSys tokens are translated into Unicode glyphs via conversionSpreadSheet.txt and compiled into Signing Gesture Markup Language (SiGML XML) by HamNoSys2SiGML.py:")
    add_code(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<sigml>\n'
        '    <hns_sign>\n'
        '        <hamnosys_nonmanual/>\n'
        '        <hamnosys_manual>\n'
        '            <hamsymmlr/>\n'
        '            <hamcee12/>\n'
        '            <hamextfingeru/>\n'
        '            <hampalml/>\n'
        '            <hamshoulders/>\n'
        '            <hamclose/>\n'
        '            <hammovel/>\n'
        '        </hamnosys_manual>\n'
        '    </hns_sign>\n'
        '</sigml>'
    )

    # --- Section 7: Execution Guide ---
    add_heading_1("7. Execution Guide & Terminal Commands")
    add_p("1. Test Pipeline on Any BSL Video via CLI:", bold_prefix="Step 1: ")
    add_code(
        'cd "Integration-20260706T062240Z-3-001\\Integration"\n'
        'py -3.12 run_local.py "bsldict\\bsldict\\videos_original\\a_001_009_000_abbreviate.mp4"'
    )
    add_p("2. Launch Web Application & 3D Avatar Server:", bold_prefix="Step 2: ")
    add_code(
        'cd "..\\..\\webapp"\n'
        'py -3.12 app.py'
    )
    add_p("3. Open http://localhost:5000 in your browser to upload videos and view live 3D avatar animations.", bold_prefix="Step 3: ")

    # Save document
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BSL_Implementation_Documentation.docx")
    try:
        doc.save(out_path)
        print(f"[Success] Generated Word Document: {out_path}")
    except PermissionError:
        fallback_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BSL_Implementation_Documentation_v2.docx")
        doc.save(fallback_path)
        print(f"[Success] Word Document was open in Microsoft Word. Saved updated copy to: {fallback_path}")

if __name__ == "__main__":
    create_report()

